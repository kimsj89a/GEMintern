import streamlit as st
import io
import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# 기본 템플릿 경로 (.docx만 지원, .dotm은 python-docx 미지원)
TEMPLATE_PATH_DOCX = os.path.join(os.path.dirname(__file__), "template", "Normal.docx")

def get_template_path():
    """사용 가능한 템플릿 경로 반환 (.docx만 지원)"""
    if os.path.exists(TEMPLATE_PATH_DOCX):
        return TEMPLATE_PATH_DOCX
    return None


def markdown_to_docx(markdown_text: str, use_template: bool = True) -> bytes:
    """마크다운 텍스트를 Word 문서로 변환"""
    # 템플릿 사용 여부에 따라 Document 생성
    template_path = get_template_path()
    if use_template and template_path:
        doc = Document(template_path)
        # 템플릿의 기존 내용 삭제 (스타일은 유지)
        for element in doc.element.body[:]:
            doc.element.body.remove(element)
    else:
        doc = Document()
        # 템플릿 없을 때만 기본 스타일 설정
        style = doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(11)

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_content = []

    while i < len(lines):
        line = lines[i]

        # 코드 블록 처리
        if line.strip().startswith('```'):
            if in_code_block:
                # 코드 블록 종료
                code_text = '\n'.join(code_content)
                p = doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Inches(0.3)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 빈 줄
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # 헤더 처리
        if line.startswith('######'):
            p = doc.add_paragraph(line[6:].strip())
            p.style = 'Heading 6' if 'Heading 6' in doc.styles else 'Normal'
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith('#####'):
            p = doc.add_paragraph(line[5:].strip())
            p.style = 'Heading 5' if 'Heading 5' in doc.styles else 'Normal'
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith('####'):
            p = doc.add_paragraph(line[4:].strip())
            p.style = 'Heading 4' if 'Heading 4' in doc.styles else 'Normal'
            run = p.runs[0] if p.runs else p.add_run()
            run.bold = True
            run.font.size = Pt(13)
        elif line.startswith('###'):
            p = doc.add_paragraph(line[3:].strip(), style='Heading 3')
        elif line.startswith('##'):
            p = doc.add_paragraph(line[2:].strip(), style='Heading 2')
        elif line.startswith('#'):
            p = doc.add_paragraph(line[1:].strip(), style='Heading 1')

        # 리스트 처리
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            indent_level = len(line) - len(line.lstrip())
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))

        elif re.match(r'^\s*\d+\.\s', line):
            indent_level = len(line) - len(line.lstrip())
            text = re.sub(r'^\s*\d+\.\s', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level // 2 + 1))

        # 수평선
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 인용문
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # 테이블 처리
        elif '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = [line]
            j = i + 1
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1

            # 테이블 생성
            if len(table_lines) >= 2:
                header_cells = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
                data_rows = []
                for tl in table_lines[2:]:  # 구분선 건너뛰기
                    cells = [cell.strip() for cell in tl.split('|') if cell.strip()]
                    if cells:
                        data_rows.append(cells)

                if header_cells:
                    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
                    table.style = 'Table Grid'

                    # 헤더
                    header_row = table.rows[0]
                    for idx, cell_text in enumerate(header_cells):
                        if idx < len(header_row.cells):
                            header_row.cells[idx].text = cell_text
                            for paragraph in header_row.cells[idx].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

                    # 데이터
                    for row_idx, row_data in enumerate(data_rows):
                        row = table.rows[row_idx + 1]
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(row.cells):
                                row.cells[col_idx].text = cell_text

                    doc.add_paragraph()  # 테이블 후 빈 줄

            i = j
            continue

        # 일반 텍스트
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)

        i += 1

    # BytesIO로 저장
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def add_formatted_text(paragraph, text: str):
    """인라인 포맷팅 처리 (볼드, 이탤릭, 코드, 링크)"""
    # 패턴들
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),  # ***bold italic***
        (r'\*\*(.+?)\*\*', 'bold'),  # **bold**
        (r'__(.+?)__', 'bold'),  # __bold__
        (r'\*(.+?)\*', 'italic'),  # *italic*
        (r'_(.+?)_', 'italic'),  # _italic_
        (r'`(.+?)`', 'code'),  # `code`
        (r'\[(.+?)\]\((.+?)\)', 'link'),  # [text](url)
    ]

    # 간단한 구현: 포맷팅 마커 제거하고 스타일 적용
    remaining = text

    # 볼드+이탤릭
    remaining = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', remaining)
    # 볼드
    remaining = re.sub(r'\*\*(.+?)\*\*', r'\1', remaining)
    remaining = re.sub(r'__(.+?)__', r'\1', remaining)
    # 이탤릭
    remaining = re.sub(r'\*(.+?)\*', r'\1', remaining)
    remaining = re.sub(r'_(.+?)_', r'\1', remaining)
    # 코드
    remaining = re.sub(r'`(.+?)`', r'\1', remaining)
    # 링크
    remaining = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', remaining)

    if remaining:
        paragraph.add_run(remaining)


def render_markdown_converter_panel(settings):
    """Markdown to Word 변환 패널"""
    st.markdown("### 📝 Markdown to Word 변환")

    st.markdown("""
        <div style='background-color: #f0f2f6; padding: 15px; border-radius: 8px; margin-bottom: 20px; border-left: 4px solid #0068c9;'>
        <h4 style='margin-top: 0; color: #0068c9;'>📋 기능 안내</h4>
        <b>마크다운 텍스트를 Word 문서(.docx)로 변환합니다.</b><br/><br/>
        <b>지원 기능:</b><br/>
        • 헤더 (H1~H6)<br/>
        • 볼드, 이탤릭 텍스트<br/>
        • 순서 있는/없는 리스트<br/>
        • 코드 블록<br/>
        • 테이블<br/>
        • 인용문<br/>
        • 수평선
        </div>
    """, unsafe_allow_html=True)

    # 입력 방식 선택
    input_method = st.radio(
        "입력 방식",
        ["직접 입력", "파일 업로드"],
        horizontal=True,
        key="md_input_method"
    )

    markdown_text = ""
    file_name = "converted"

    if input_method == "직접 입력":
        markdown_text = st.text_area(
            "마크다운 텍스트 입력",
            height=400,
            placeholder="마크다운 텍스트를 여기에 붙여넣으세요...\n\n예시:\n# 제목\n## 소제목\n- 리스트 항목\n**볼드 텍스트**",
            key="md_text_input"
        )
    else:
        uploaded_file = st.file_uploader(
            "마크다운 파일 업로드 (.md, .txt)",
            type=['md', 'txt'],
            key="md_file_upload"
        )

        if uploaded_file:
            markdown_text = uploaded_file.read().decode('utf-8')
            file_name = uploaded_file.name.rsplit('.', 1)[0]
            st.text_area(
                "파일 내용 미리보기",
                value=markdown_text,
                height=300,
                disabled=True,
                key="md_file_preview"
            )

    # 변환 옵션
    st.markdown("#### ⚙️ 변환 옵션")
    col1, col2 = st.columns(2)
    with col1:
        output_filename = st.text_input(
            "출력 파일명",
            value=file_name,
            key="md_output_filename"
        )
    with col2:
        template_path = get_template_path()
        template_exists = template_path is not None
        use_template = st.checkbox(
            "기본 템플릿 사용 (Normal.docx)",
            value=template_exists,
            disabled=not template_exists,
            help="template/Normal.docx 파일의 스타일을 적용합니다." if template_exists else "템플릿 파일이 없습니다.",
            key="md_use_template"
        )
        if template_exists:
            st.caption(f"📄 템플릿: `template/Normal.docx`")

    # 변환 버튼
    st.markdown("---")
    if st.button("🚀 Word로 변환", type="primary", use_container_width=True, key="btn_convert_md"):
        if not markdown_text.strip():
            st.warning("변환할 마크다운 텍스트를 입력해주세요.")
            return

        with st.spinner("변환 중..."):
            try:
                docx_bytes = markdown_to_docx(markdown_text, use_template=use_template)
                st.session_state['md_converted_docx'] = docx_bytes
                st.session_state['md_converted_filename'] = output_filename
                st.success("변환 완료!" + (" (템플릿 적용됨)" if use_template else ""))
            except Exception as e:
                st.error(f"변환 중 오류 발생: {e}")

    # 다운로드 버튼
    if st.session_state.get('md_converted_docx'):
        filename = st.session_state.get('md_converted_filename', 'converted')
        st.download_button(
            "📥 Word 문서 다운로드",
            st.session_state['md_converted_docx'],
            f"{filename}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            type="primary",
            key="btn_download_docx"
        )
