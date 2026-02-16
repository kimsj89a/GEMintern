"""
Markdown to Word conversion utilities.
Uses pypandoc (Pandoc) for high-quality conversion with reference-doc template.
"""
import io
import os
import tempfile

# 기본 템플릿 경로
TEMPLATE_PATH_DOCX = os.path.join(os.path.dirname(__file__), "template", "Normal.docx")


def get_template_path():
    """사용 가능한 템플릿 경로 반환 (.docx만 지원)"""
    if os.path.exists(TEMPLATE_PATH_DOCX):
        return TEMPLATE_PATH_DOCX
    return None


def markdown_to_docx(markdown_text: str, use_template: bool = True) -> bytes:
    """마크다운 텍스트를 Word 문서로 변환 (pypandoc 사용)"""
    import pypandoc

    template_path = get_template_path() if use_template else None

    extra_args = [
        '--toc',
        '--toc-depth=3',
        '--wrap=none',
    ]
    if template_path:
        extra_args.append(f'--reference-doc={template_path}')

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = tmp.name

    try:
        pypandoc.convert_text(
            markdown_text,
            'docx',
            format='markdown+pipe_tables+footnotes+hard_line_breaks+strikeout',
            outputfile=tmp_path,
            extra_args=extra_args,
        )

        with open(tmp_path, 'rb') as f:
            docx_bytes = f.read()

        # 한글 폰트 후처리 (reference-doc에 설정 안 된 경우 대비)
        docx_bytes = _apply_korean_font(docx_bytes)

        return docx_bytes
    finally:
        if os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except OSError:
                pass


def _apply_korean_font(docx_bytes: bytes) -> bytes:
    """python-docx로 한글 폰트 후처리"""
    try:
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(io.BytesIO(docx_bytes))

        font_name = '맑은 고딕'

        # Normal 스타일에 한글 폰트 설정
        style = doc.styles['Normal']
        style.font.name = font_name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 모든 paragraph의 run에 한글 폰트 적용
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        # 테이블 셀 내 텍스트에도 적용
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = font_name
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception:
        # 후처리 실패 시 원본 반환
        return docx_bytes
